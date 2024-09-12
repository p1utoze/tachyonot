import PyPDF2
from textsplitter import TextSplitter
from llama_cpp import Llama
import os
from docx import Document
from ..utils.schema import Document
from chardet import detect


class LLamaCPPEmbedding:

    def __init__(
        self,
        embedding_model_path: str,
        chunk_size: int = 1024,
        chunk_overlap: int = 100,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_model = Llama(
            model_path=embedding_model_path, embedding=True, verbose=False
        )
        self.text_splitter = TextSplitter(
            max_token_size=self.chunk_size,
            end_sentence=True,
            preserve_formatting=True,
            language="english",
        )

    def embed_text(self, document: str):
        """
        Embed the input text using the specified embedding model.
        :param document: The input text to be embedded
        :return: A list of Document objects, each containing a text chunk and its corresponding embedding
        """

        chunks = self._split_text(document)
        embeddings = self.embedding_model.create_embedding(chunks)
        embed_text = []
        for chunk, embedding in zip(chunks, embeddings["data"]):
            embed_text.append(Document(text=chunk, embedding=embedding["embedding"]))
        return embed_text

    def embed_prompt(self, prompt: str):
        """
        Embed the input prompt using the specified embedding model.
        :param prompt: The input prompt to be embedded
        :return: The embedding of the input prompt
        """
        return self.embedding_model.create_embedding(prompt)["data"][0]["embedding"]

    def embed_file(self, file_path):
        """
        Embed the input file using the specified embedding model.
        :param file_path: The path to the input file
        :return: A list of Document objects, each containing a text chunk and its corresponding embedding
        """
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()

        if file_extension == ".txt":
            text = self._read_text_file(file_path)
        elif file_extension == ".pdf":
            text = self._read_pdf_file(file_path)
        elif file_extension == ".docx":
            text = self._read_docx_file(file_path)
        else:
            raise "Unsupported file type"

        return self.embed_text(text)

    def embed_documents(self, directory_path: str):
        """
        Embed all the documents in the specified directory using the specified embedding model.
        :param directory_path: The path to the directory containing the documents
        :return: A list of Document objects, each containing a text chunk and its corresponding embedding
        """
        document_paths = [
            os.path.join(directory_path, doc_name)
            for doc_name in os.listdir(directory_path)
        ]
        embed_text = []
        for document_path in document_paths:
            embed_text.extend(self.embed_file(document_path))
        return embed_text

    def _split_text(self, text):
        """
        Split the input text into chunks using TextSplitter.
        :param text: The input text to be split
        :return: A list of text chunks
        """
        chunks = self.text_splitter.split_text(text)

        # Add overlap
        overlapped_chunks = []
        for i in range(len(chunks)):
            if i > 0:
                overlap = chunks[i - 1][-self.chunk_overlap :]
                overlapped_chunks.append(overlap + chunks[i])
            else:
                overlapped_chunks.append(chunks[i])

        return overlapped_chunks

    def _read_text_file(self, file_path):
        with open(file_path, "rb") as file:
            text = file.read().decode(detect(file.read())["encoding"])
        return text

    def _read_pdf_file(self, file_path):
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

    def _read_docx_file(self, file_path):
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text


if __name__ == "__main__":
    root = "data"
    embeder = LLamaCPPEmbedding(
        "model/all-MiniLM-L6-v2.F32.gguf", chunk_size=100, chunk_overlap=20
    )
    print(len(embeder.embed_documents("data")))
